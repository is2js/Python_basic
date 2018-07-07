from employee import Employee
import csv
from docx import Document


f = open('employee_list.csv', 'r')
# lines = f.readlines()
lines = csv.reader(f)

employee_list = list()
i=0
for line in lines:
    if i>0:
        # csv모듈 안쓸 때 코드
        # info = line.split(',')
        # employee_list.append( Employee( info[0], int(info[1]), int(info[2]), int(info[3]) ))
        employee_list.append( Employee( line[0], int(line[1]), int(line[2]), int(line[3])) )
    i+=1
f.close()

csv_file = open('result.csv', 'w', newline='')
csv_writer = csv.writer(csv_file)

# csv_file.write('이름, 출근시간, 퇴근시간, 시급, 근무시간, 일당\n')
csv_writer.writerow([ '이름', '출근시간', '퇴근시간', '시급', '근무시간', '일당'])
for e in employee_list:
    # csv모듈 안쓸 때 코드
    # csv_file.write('{}, {}, {}, {}, {}, {}\n'.format(
    #     e.name, e.work_start, e.work_finish, e.wage_per_hour, e.worked_hours(), e.wage_of_the_day()))
    csv_writer.writerow([
        e.name, e.work_start, e.work_finish, e.wage_per_hour, e.worked_hours(), e.wage_of_the_day()
    ])
csv_file.close()


#문서화 코드---------------------------------------------------------------------
#양식을 넣은 도큐먼트 객체 생성
document = Document('form.docx')

#문서 전체 제목
document.add_heading('만두가게 임금 산정결과 문서', 0)

#제목1수준으로 각 일꾼들의 기본정보를 .format으로 문장형으로 생성
document.add_heading('오늘 출근직원의 노동 정보', 1)
#원형리스트형식으로 각 인스턴스 속 정보 뽑아내기
for employee in employee_list:
    document.add_paragraph('{}님은 {}시 출근하여 {}시 퇴근하였습니다.'.format(
        employee.name, employee.work_start, employee.work_finish
    ) , style='circle')


#표 생성(한행씩 / 들어갈 칼럼 갯수 명시)
table = document.add_table(rows=1, cols=6)
#테이블 스타일 지정(문서에서 미리 생성한 사용자 지정 양식)
table.style = 'table'

#첫ROW는 직접 적어준다. 칼럼명이기 때문.
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '이름'
hdr_cells[1].text = '출근'
hdr_cells[2].text = '퇴근'
hdr_cells[3].text = '근무시간'
hdr_cells[4].text = '시급'
hdr_cells[5].text = '일당'

#2번째 ROW부터는 정보를 대입해야 하므로 인스턴스 관리 리스트를 이용한 for문으로 각 row 채우기
#각 셀들에 접근.text 만 가능하며, str()로 문자열로 반드시 바꿔서대입해야한다.int안됨.
for employee in employee_list:
    row_cells = table.add_row().cells
    row_cells[0].text = str(employee.name)
    row_cells[1].text = str(employee.work_start)
    row_cells[2].text = str(employee.work_finish)
    row_cells[3].text = str(employee.worked_hours())
    row_cells[4].text = str(int(employee.wage_per_hour))
    row_cells[5].text = str(int(employee.wage_of_the_day()))

# docx파일 생성을 위한 save('파일명')
document.save('result.docx')