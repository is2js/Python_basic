from employee import Employee
import csv
from docx import Document


f = open('employee_list.csv', 'r')
lines = csv.reader(f)
#for문에 들어가기전에, 모든 인스턴스들을 관리할 리스트 생성
employee_list = list()
#for문에 들어가기전에, 첫 줄에서는 작업안하도록 매개변수 생성
i = 0
for line in lines:
    #첫줄에서는 작업안하도록 if문 걸어주기
    if i > 0:
        employee_list.append( Employee(line[0], int( line[1] ), int( line[2] ), int( line[3] )))
    #첫줄 작업안하도록 1단계식 띄워넘기
    i += 1
f.close()

#list를 확인할 때는, len()와, for문을 돌면서 각 요소들 확인
print( len(employee_list) )
for emp in employee_list:
    print(emp.name, emp.work_start, emp.work_finish, emp.worked_hours(), emp.wage_of_the_day())


#세부적으로 확인하는 for문을 그대로 활용해서, 파일 쓰기를 한다.
f_output = open('employee_result.csv', 'w', newline='')
csv_writer = csv.writer(f_output)
csv_writer.writerow(['이름', '출근시간', '퇴근시간', '시급', '근무시간', '일당'])
#확인 복붙 부분
for emp in employee_list:
    csv_writer.writerow([
        emp.name, emp.work_start, emp.work_finish, emp.wage_per_hour, emp.worked_hours(), emp.wage_of_the_day()
    ])
    # f_output.write( '{}, {}, {}, {}, {}, {}\n'.format(
    #     emp.name, emp.work_start, emp.work_finish, emp.wage_per_hour,emp.worked_hours(), emp.wage_of_the_day()))

f_output.close()

#문서생성
document = Document('form.docx')

#제목
document.add_heading('만두가게 임금 산정 결과',0)
#제목1
document.add_heading('오늘 출근한 직원', 1)
for employee in employee_list:
    document.add_paragraph('{}님은 {}시 출근 ~ {}시 퇴근하였습니다.'.format(
        employee.name, employee.work_start, employee.work_finish))



#제목2
document.add_heading('출근직원의 시급과 일당', 2)
#테이블 생성
table = document.add_table(rows=1, cols=6)
#테이블 스타일 지정
table.style = 'table'
#첫ROW만, 직접 적어준다. 특이하게, 각 첫행의 각 셀들을 리스트변수로 받는다.
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '이름'
hdr_cells[1].text = '출근'
hdr_cells[2].text = '퇴근'
hdr_cells[3].text = '근무시간'
hdr_cells[4].text = '시급'
hdr_cells[5].text = '일당'

#인스턴스 관리 리스트를 이용한 for문으로 각 row 채우기
#각 셀들에 접근.text 만 가능하며, str()로 문자열로 반드시 바꿔서대입해야한다.int안됨.
for employee in employee_list:
    row_cells = table.add_row().cells
    row_cells[0].text = str(employee.name)
    row_cells[1].text = str(employee.work_start)
    row_cells[2].text = str(employee.work_finish)
    row_cells[3].text = str(employee.worked_hours())
    row_cells[4].text = str(int(employee.wage_per_hour))
    row_cells[5].text = str(int(employee.wage_of_the_day()))

#docx파일 생성을 위한 save('파일명')
document.save('mandoo_wage.docx')
