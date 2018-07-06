from docx import Document

#사용하기 위한 변수 선언
document = Document('form.docx')

#제목을 위한 add_heading( '제목', 0)함수 사용
document.add_heading('파이썬으로 워드파일 만들기', 0)

document.add_heading('제목1양식-1 워드파일', 1)
document.add_heading('제목2양식 워드파일', 2)
document.add_heading('제목3양식 워드파일', 3)

#본문(paragraph)
document.add_paragraph('동해물과 백두산이 마르고 닳도록 하느님이 보우하사 우리나라 만세.')

#새로운 스타일(동그라미스타일)도 본문(paragraph) + style인자.
document.add_paragraph('긴급속보! 조재성, 김석영 결혼', style='circle')
document.add_paragraph('긴급속보! 조재성 한의대 졸업!', style='circle')

document.add_paragraph('긴급속보! 조재성 한의대 졸업!', style='check')
document.add_paragraph('긴급속보! 조재성 한의대 졸업!', style='check')







#테이블 생성
table = document.add_table(rows=1, cols=3)
#테이블 스타일 지정
table.style = 'table'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'

row_cells = table.add_row().cells
row_cells[0].text = str(123)
row_cells[1].text = str(123)
row_cells[2].text = '아무거나 넣어보자.'

#docx파일 생성을 위한 save('파일명')
document.save('demo.docx')


